"""
Resume parsing service for extracting text from PDF and DOCX files.

Supports:
- PDF files using PyPDF2
- DOCX files using python-docx
- Plain text files
"""

import io
import logging
from typing import Tuple
from pathlib import Path

from pypdf import PdfReader
from pypdf.errors import PdfReadError
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)


# ============================================
# Supported File Types
# ============================================

SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


# ============================================
# PDF Extraction
# ============================================

def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file content.

    Args:
        file_content: Raw bytes of the PDF file

    Returns:
        Extracted text from all pages

    Raises:
        ValueError: If PDF is encrypted or cannot be read
    """
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)

        # Check if encrypted
        if reader.is_encrypted:
            try:
                # Try empty password first
                reader.decrypt("")
            except Exception:
                raise ValueError(
                    "PDF is encrypted and cannot be read. "
                    "Please provide an unencrypted PDF."
                )

        # Extract text from all pages
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num}: {e}")
                continue

        if not text_parts:
            raise ValueError(
                "No text could be extracted from the PDF. "
                "The file may be image-based or corrupted."
            )

        return "\n\n".join(text_parts)

    except PdfReadError as e:
        raise ValueError(f"Invalid or corrupted PDF file: {e}")
    except Exception as e:
        if "encrypted" in str(e).lower():
            raise ValueError("PDF is encrypted and cannot be read.")
        raise ValueError(f"Failed to read PDF: {e}")


# ============================================
# DOCX Extraction
# ============================================

def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file content.

    Args:
        file_content: Raw bytes of the DOCX file

    Returns:
        Extracted text preserving paragraph structure

    Raises:
        ValueError: If DOCX cannot be read
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # Extract tables (often contain skills/experience)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        if not text_parts:
            raise ValueError(
                "No text could be extracted from the DOCX file. "
                "The file may be empty or corrupted."
            )

        return "\n\n".join(text_parts)

    except PackageNotFoundError:
        raise ValueError("Invalid or corrupted DOCX file.")
    except Exception as e:
        raise ValueError(f"Failed to read DOCX file: {e}")


# ============================================
# Plain Text Extraction
# ============================================

def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from plain text file content.

    Args:
        file_content: Raw bytes of the text file

    Returns:
        Decoded text content

    Raises:
        ValueError: If text cannot be decoded
    """
    # Try common encodings
    encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

    for encoding in encodings:
        try:
            return file_content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError(
        "Could not decode text file. "
        "Please ensure the file is in UTF-8 or a common text encoding."
    )


# ============================================
# Main Parser Function
# ============================================

def parse_resume(
    file_content: bytes,
    filename: str,
    content_type: str = None
) -> Tuple[str, str]:
    """
    Parse resume file and extract text.

    Args:
        file_content: Raw bytes of the file
        filename: Original filename
        content_type: MIME type of the file (optional)

    Returns:
        Tuple of (extracted_text, file_type)

    Raises:
        ValueError: If file type not supported or extraction fails
    """
    # Validate file size
    if len(file_content) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File size exceeds maximum of {MAX_FILE_SIZE_MB}MB. "
            "Please upload a smaller file."
        )

    if len(file_content) == 0:
        raise ValueError("File is empty.")

    # Determine file type
    file_ext = Path(filename).suffix.lower()

    if file_ext not in SUPPORTED_EXTENSIONS:
        # Try to use content type
        if content_type and content_type in SUPPORTED_MIME_TYPES:
            file_type = SUPPORTED_MIME_TYPES[content_type]
        else:
            raise ValueError(
                f"Unsupported file type: {file_ext}. "
                f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
            )
    else:
        file_type = file_ext.lstrip(".")

    # Extract text based on file type
    logger.info(f"Parsing {file_type.upper()} file: {filename}")

    if file_type == "pdf":
        text = extract_text_from_pdf(file_content)
    elif file_type == "docx":
        text = extract_text_from_docx(file_content)
    elif file_type == "txt":
        text = extract_text_from_txt(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    logger.info(f"Extracted {len(text)} characters from {filename}")

    return text, file_type


# ============================================
# Validation Functions
# ============================================

def validate_file_type(filename: str, content_type: str = None) -> bool:
    """
    Check if file type is supported.

    Args:
        filename: Original filename
        content_type: MIME type (optional)

    Returns:
        True if file type is supported
    """
    file_ext = Path(filename).suffix.lower()

    if file_ext in SUPPORTED_EXTENSIONS:
        return True

    if content_type and content_type in SUPPORTED_MIME_TYPES:
        return True

    return False


def get_file_type(filename: str, content_type: str = None) -> str:
    """
    Get normalized file type from filename or content type.

    Args:
        filename: Original filename
        content_type: MIME type (optional)

    Returns:
        File type string (pdf, docx, txt)

    Raises:
        ValueError: If file type not supported
    """
    file_ext = Path(filename).suffix.lower()

    if file_ext in SUPPORTED_EXTENSIONS:
        return file_ext.lstrip(".")

    if content_type and content_type in SUPPORTED_MIME_TYPES:
        return SUPPORTED_MIME_TYPES[content_type]

    raise ValueError(f"Unsupported file type: {file_ext or content_type}")
